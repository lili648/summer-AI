"""
========================================
  RAG 问答系统 — Streamlit + LangChain + ChromaDB
  基于上传文档的智能问答（检索增强生成）
  API: 智谱 (ZhipuAI)
========================================
"""
import os
import sys
import tempfile
from pathlib import Path
from typing import List

import streamlit as st

# ── LangChain ──────────────────────────────────────────────
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings

# Windows 控制台 UTF-8
sys.stdout.reconfigure(encoding="utf-8")

# ══════════════════════════════════════════════════════════
#  常量 & 配置
# ══════════════════════════════════════════════════════════

PERSIST_DIR = Path(__file__).parent / "chroma_db"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 100

ZHIPU_BASE_URL = "https://open.bigmodel.cn/api/paas/v4/"
ZHIPU_EMBEDDING_MODEL = "embedding-2"

# ══════════════════════════════════════════════════════════
#  自定义智谱 Embedding（直接调用 API，无需 tiktoken）
# ══════════════════════════════════════════════════════════


class ZhipuEmbeddings(Embeddings):
    """通过智谱 OpenAI 兼容接口调用 embedding，不依赖 tiktoken"""

    def __init__(self, api_key: str, model: str = ZHIPU_EMBEDDING_MODEL):
        from openai import OpenAI

        self._client = OpenAI(api_key=api_key, base_url=ZHIPU_BASE_URL)
        self.model = model

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """批量 embedding"""
        resp = self._client.embeddings.create(model=self.model, input=texts)
        return [item.embedding for item in resp.data]

    def embed_query(self, text: str) -> List[float]:
        """单条 embedding"""
        resp = self._client.embeddings.create(model=self.model, input=text)
        return resp.data[0].embedding


# ══════════════════════════════════════════════════════════
#  文档加载器（每种格式返回纯文本）
# ══════════════════════════════════════════════════════════


def load_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def load_md(file_path: str) -> str:
    return load_txt(file_path)


def load_pdf(file_path: str) -> str:
    """用 pypdf 解析 PDF"""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def load_docx(file_path: str) -> str:
    """加载 .docx 文件"""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        st.error("请安装 python-docx: pip install python-docx")
        return ""


def extract_text_from_file(file_path: str) -> str:
    """根据后缀名自动选择加载器"""
    ext = Path(file_path).suffix.lower()
    loaders = {
        ".txt": load_txt,
        ".md": load_md,
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".doc": load_docx,
    }
    loader = loaders.get(ext)
    if loader is None:
        return ""
    return loader(file_path)


# ══════════════════════════════════════════════════════════
#  文本分块
# ══════════════════════════════════════════════════════════


def split_text(text: str, source_name: str = "") -> List[Document]:
    """将文本切分为 Document 块"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", ".", " ", ""],
        length_function=len,
    )
    chunks = splitter.split_text(text)
    docs = []
    for i, chunk in enumerate(chunks):
        m = {
            "source": source_name,
            "chunk": i,
            "total": len(chunks),
        }
        docs.append(Document(page_content=chunk, metadata=m))
    return docs


# ══════════════════════════════════════════════════════════
#  ChromaDB 向量存储
# ══════════════════════════════════════════════════════════


@st.cache_resource(show_spinner="连接向量数据库...")
def get_vector_store(_embeddings, persist_dir: str = str(PERSIST_DIR)):
    """获取或创建 Chroma 向量库（持久化到磁盘）"""
    persist_path = Path(persist_dir)
    persist_path.mkdir(parents=True, exist_ok=True)
    vector_store = Chroma(
        collection_name="rag_docs",
        embedding_function=_embeddings,
        persist_directory=str(persist_path),
    )
    return vector_store


def add_documents_to_store(vector_store: Chroma, docs: List[Document]):
    """将文档块写入向量库"""
    vector_store.add_documents(docs)


@st.cache_data(ttl=10, show_spinner=False)
def get_collection_stats(_vector_store) -> dict:
    """获取集合中文档数量（带缓存）"""
    try:
        count = _vector_store._collection.count()
    except Exception:
        count = 0
    return {"count": count}


# ══════════════════════════════════════════════════════════
#  检索 + 生成
# ══════════════════════════════════════════════════════════


def build_context(vector_store: Chroma, query: str, k: int = 4) -> str:
    """检索最相关的文档块，拼接成上下文"""
    results = vector_store.similarity_search(query, k=k)
    if not results:
        return ""
    segments = []
    for i, doc in enumerate(results):
        src = doc.metadata.get("source", "未知")
        segments.append(f"【片段 {i + 1}】（来源：{src}）\n{doc.page_content}")
    return "\n\n".join(segments)


def query_llm(
    client,
    model: str,
    system_prompt: str,
    user_query: str,
    context: str,
    history: list,
):
    """向智谱 LLM 发送带 RAG 上下文的请求（流式输出）"""
    if context.strip():
        rag_system = (
            f"{system_prompt}\n\n"
            "## 上下文资料（基于你已上传的文档）\n"
            "请优先根据以下资料回答问题。如果资料不足以回答，"
            "请说明「根据已有资料无法完整回答」，然后补充你的理解。\n"
            "不要编造资料中没有的信息。\n\n"
            f"{context}"
        )
    else:
        rag_system = system_prompt

    messages = [{"role": "system", "content": rag_system}]
    messages += history
    messages.append({"role": "user", "content": user_query})

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        stream=True,
    )
    return response


# ══════════════════════════════════════════════════════════
#  Streamlit UI
# ══════════════════════════════════════════════════════════

st.set_page_config(
    page_title="RAG 问答系统",
    page_icon="📚",
    layout="wide",
)

# ── 侧边栏：API & 模型设置 ───────────────────────────────

with st.sidebar:
    st.header("🔑 API 设置")
    api_key = os.environ.get("ZHIPU_API_KEY") or st.text_input(
        "智谱 API Key",
        type="password",
        placeholder="xxx.yyy",
        help="优先读取 ZHIPU_API_KEY 环境变量",
    )
    if not api_key:
        st.warning("⚠️ 请设置 ZHIPU_API_KEY 或在下方填入 API Key")
        st.stop()

    # 初始化智谱客户端（LLM + Embedding 共用）
    from openai import OpenAI

    llm_client = OpenAI(api_key=api_key, base_url=ZHIPU_BASE_URL)
    embeddings = ZhipuEmbeddings(api_key=api_key)

    st.divider()
    st.header("⚙️ 模型设置")
    model = st.selectbox(
        "对话模型",
        ["glm-4-flash", "glm-4-plus", "glm-4-air", "glm-4-long"],
        index=0,
        help="glm-4-flash: 快速免费 | glm-4-plus: 效果最佳 | glm-4-air: 均衡",
    )
    k_retrieval = st.slider(
        "检索文档块数 (k)",
        min_value=1,
        max_value=10,
        value=4,
        help="根据问题检索多少个相关文档片段",
    )
    system_prompt = st.text_area(
        "系统提示词",
        value="你是一个专业的文档问答助手，请基于提供的资料准确回答用户问题。"
        "回答时请注明信息来源。如果资料不足，请明确告知。",
        height=100,
    )

    st.divider()
    if st.button("🧹 清空对话", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

# ── 主界面 ───────────────────────────────────────────────

st.title("📚 RAG 智能问答系统")
st.caption("上传文档 → 自动建立知识库 → 基于文档内容回答问题")

vector_store = get_vector_store(embeddings)

# 会话状态
if "messages" not in st.session_state:
    st.session_state.messages = []
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = set()

# ── 两个 Tab：文档管理 / 问答 ───────────────────────────

tab_doc, tab_chat = st.tabs(["📄 文档管理", "💬 问答"])

# =================== Tab 1: 文档管理 =====================

with tab_doc:
    col_left, col_right = st.columns([1, 1])

    with col_left:
        st.subheader("📤 上传文档")
        uploaded_files = st.file_uploader(
            "支持 PDF / TXT / MD / DOCX",
            type=["pdf", "txt", "md", "docx"],
            accept_multiple_files=True,
        )

        if uploaded_files:
            with st.status("⏳ 正在处理文档...", expanded=True) as status:
                new_count = 0
                for up_file in uploaded_files:
                    file_key = f"{up_file.name}:{up_file.size}"
                    if file_key in st.session_state.uploaded_files:
                        st.write(f"⏭️ 跳过已处理的 `{up_file.name}`")
                        continue

                    # 保存到临时文件
                    suffix = Path(up_file.name).suffix
                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=suffix
                    ) as tmp:
                        tmp.write(up_file.getbuffer())
                        tmp_path = tmp.name

                    try:
                        st.write(f"📖 读取 `{up_file.name}` ...")
                        text = extract_text_from_file(tmp_path)
                        if not text.strip():
                            st.warning(f"⚠️ `{up_file.name}` 内容为空或无法解析")
                            os.unlink(tmp_path)
                            continue

                        st.write(f"✂️ 分块中 ...")
                        docs = split_text(text, source_name=up_file.name)
                        st.write(f"➕ 写入向量库（{len(docs)} 块）...")
                        add_documents_to_store(vector_store, docs)
                        st.session_state.uploaded_files.add(file_key)
                        new_count += 1
                    except Exception as e:
                        st.error(f"❌ `{up_file.name}` 处理失败：{e}")
                    finally:
                        os.unlink(tmp_path)

                status.update(
                    label=f"✅ 完成！新增了 {new_count} 个文件到知识库",
                    state="complete",
                )
            # 清除统计缓存，刷新显示
            get_collection_stats.clear()

    with col_right:
        st.subheader("📊 知识库状态")
        stats = get_collection_stats(vector_store)
        doc_count = stats["count"]
        if doc_count > 0:
            st.metric("文档片段总数", doc_count)
            st.success(f"知识库中有 **{doc_count}** 个文本片段可供检索")
            if st.button("🗑️ 清空全部知识库", type="secondary"):
                try:
                    vector_store._collection.delete()
                    get_collection_stats.clear()
                    st.session_state.uploaded_files = set()
                    st.rerun()
                except Exception as e:
                    st.error(f"清空失败：{e}")
        else:
            st.info("💡 知识库为空，请上传文档")

    st.divider()
    st.caption("💡 提示：上传后即可切换到「问答」Tab 开始提问")

# =================== Tab 2: 问答 =========================

with tab_chat:
    stats = get_collection_stats(vector_store)
    if stats["count"] == 0:
        st.warning("📭 知识库为空，请先在「文档管理」页面上传文档")
    else:
        st.success(f"✅ 知识库已就绪（{stats['count']} 个片段）")

    # 回放历史消息
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            if msg.get("context"):
                with st.expander("📎 检索到的资料", expanded=False):
                    st.markdown(msg["context"])
            if msg.get("reasoning"):
                with st.expander("🧠 思考过程"):
                    st.markdown(msg["reasoning"])
            st.markdown(msg["content"])

    # 聊天输入
    prompt = st.chat_input("请输入你的问题...")

    if prompt:
        # 显示用户消息
        st.chat_message("user").markdown(prompt)

        # 检索引擎
        with st.status("🔍 正在检索知识库...", expanded=False) as search_status:
            context = build_context(vector_store, prompt, k=k_retrieval)
            if context:
                search_status.update(
                    label="✅ 已检索到相关文档片段",
                    state="complete",
                )
            else:
                search_status.update(
                    label="⚠️ 未检索到相关文档内容",
                    state="error",
                )

        # 组装历史（不含 context 和 reasoning 的纯消息）
        history = [
            {"role": m["role"], "content": m["content"]}
            for m in st.session_state.messages
        ]

        # 调用 LLM
        with st.chat_message("assistant"):
            response = query_llm(
                client=llm_client,
                model=model,
                system_prompt=system_prompt,
                user_query=prompt,
                context=context,
                history=history,
            )

            reasoning_text = ""
            answer_text = ""
            reasoning_box = st.empty()
            answer_box = st.empty()

            for chunk in response:
                if not chunk.choices:
                    continue
                delta = chunk.choices[0].delta

                if getattr(delta, "reasoning_content", None):
                    reasoning_text += delta.reasoning_content
                    with reasoning_box.expander("🧠 思考过程", expanded=True):
                        st.markdown(reasoning_text)

                if delta.content:
                    answer_text += delta.content
                    answer_box.markdown(answer_text)

            if reasoning_text:
                with reasoning_box.expander("🧠 思考过程"):
                    st.markdown(reasoning_text)

            # 折叠显示检索到的上下文
            if context:
                with st.expander("📎 查看引用的文档片段", expanded=False):
                    st.markdown(context)

        # 保存到历史
        st.session_state.messages.append(
            {
                "role": "user",
                "content": prompt,
            }
        )
        st.session_state.messages.append(
            {
                "role": "assistant",
                "content": answer_text,
                "reasoning": reasoning_text,
                "context": context,
            }
        )
        st.rerun()